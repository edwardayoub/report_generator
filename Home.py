import streamlit as st
import pandas as pd
import numpy as np
import io
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
import zipfile
from datetime import datetime

def main():
    st.title("EAyoub Report Generator 15min")
    
    meta_file = st.file_uploader("Upload Meta File", type=["xlsx"])
    locations_file = st.file_uploader("Upload Locations Data", type=["csv"])
    
    if meta_file is not None and locations_file is not None:
        trip_meta = pd.read_excel(meta_file)
        trip_data = pd.read_csv(locations_file)
        
        location_names = [col for col in trip_data.columns if 'LOC' in col or 'NT' in col]
        
        # Create LOC_ID choices
        loc_ids = [f"LOC_{i+1}" for i in range(len(location_names))]
        selected_loc_id = st.selectbox("Display the following location", loc_ids)
        
        if st.button("Process"):
            # Process the data
            result_list = process_data(trip_data, trip_meta, location_names)
            
            # Display data table
            loc_index = int(selected_loc_id.split('_')[1]) - 1
            if loc_index < len(location_names):
                loc_name = location_names[loc_index]
                st.write(f"Data for {loc_name}")
                st.dataframe(result_list[loc_name]['dmodel'])
            else:
                st.write("Invalid LOC_ID selected.")
            
            # Provide download link
            zip_buffer = generate_zip(result_list)
            st.download_button(
                label="Download",
                data=zip_buffer.getvalue(),
                file_name=f"my_data_{datetime.now().strftime('%Y-%m-%d')}.zip",
                mime="application/zip"
            )

def process_data(trip_data, trip_meta, location_names):
    result_list = {}
    first_date = pd.to_datetime(trip_data['Date'].iloc[0]).date()
    for idx, loc_name in enumerate(location_names):
        LOC_col = loc_name
        entry_numbers = trip_meta[trip_meta['Location'] == loc_name].index.tolist()
        if not entry_numbers:
            continue
        entry_number = entry_numbers[0]
        
        # Adjusted column names here
        trip_job = trip_meta.loc[entry_number, 'Contract']
        trip_client = trip_meta.loc[entry_number, 'Client']
        trip_stationid = trip_meta.loc[entry_number, 'Facility']
        trip_location = loc_name.replace('LOC', '')
        trip_atr = trip_meta.loc[entry_number, 'ATR']
        # convert trip_atr to int
        trip_atr = int(float(trip_atr)) if pd.notna(trip_atr) else 0
        trip_site_des = trip_meta.loc[entry_number, 'Site']
        trip_site_gps = trip_meta.loc[entry_number, 'Site2']
        
        # Create dmodel DataFrame
        dmodel = create_dmodel(trip_data, LOC_col)
        
        # Calculate totals and other metrics
        totals, day_totals, splits, trip_peak, trip_peak_vol, trip_phf, peak_indices_list = calculate_metrics(dmodel)
        
        # Generate Word document
        doc = generate_word_doc(trip_job, trip_client, trip_stationid, trip_location, trip_atr,
                                trip_site_des, trip_site_gps, dmodel, totals, day_totals, splits,
                                trip_peak, trip_peak_vol, trip_phf, peak_indices_list, first_date)
        
        # Generate Excel file
        excel_file = generate_excel_file(trip_job, trip_client, trip_stationid, trip_location,
                                         trip_atr, trip_site_des, trip_site_gps, dmodel, totals,
                                         day_totals, splits, trip_peak, trip_peak_vol, trip_phf, first_date)
        
        # Store in result_list
        result_list[loc_name] = {
            'doc': doc,
            'excel': excel_file,
            'dmodel': dmodel
        }
    return result_list

def create_dmodel(trip_data, LOC_col):
    # Create the dmodel DataFrame as in R code
    dmodel = pd.DataFrame()
    dmodel['Time'] = trip_data['Time'].iloc[0:48].reset_index(drop=True)
    for i in range(7):
        AM_col = f"AM_{i+1}"
        PM_col = f"PM_{i+1}"
        start_AM = i * 96
        end_AM = start_AM + 48
        start_PM = end_AM
        end_PM = start_PM + 48
        dmodel[AM_col] = trip_data[LOC_col].iloc[start_AM:end_AM].reset_index(drop=True).round(0)
        dmodel[PM_col] = trip_data[LOC_col].iloc[start_PM:end_PM].reset_index(drop=True).round(0)
    # Calculate averages
    dmodel['AM'] = dmodel[[f"AM_{i+1}" for i in range(7)]].mean(axis=1).round(0)
    dmodel['PM'] = dmodel[[f"PM_{i+1}" for i in range(7)]].mean(axis=1).round(0)
    return dmodel

def calculate_metrics(dmodel):
    # Totals
    totals = dmodel.iloc[:,1:].sum()
    # Day totals
    day_totals = []
    for i in range(7):
        day_total = dmodel[f"AM_{i+1}"].sum() + dmodel[f"PM_{i+1}"].sum()
        day_totals.append(day_total)
    avg_total = dmodel['AM'].sum() + dmodel['PM'].sum()
    day_totals.append(avg_total)
    # Splits
    splits = (totals / np.repeat(day_totals, [2]*8)) * 100
    splits = splits.replace(np.nan, 0).round(1)
    # Trip peak calculations
    trip_peak = []
    trip_peak_vol = []
    trip_phf = []
    peak_indices_list = []
    for idx, col in enumerate(dmodel.columns[1:]):
        results = dmodel[col].values
        n = len(results)
        max_sum = -np.inf
        peak_index = None
        # Consider four different shifts
        for shift in range(4):
            shifted_results = results[shift:n - (3 - shift)]
            sums = np.array([sum(shifted_results[i:i + 4]) for i in range(len(shifted_results) - 3)])
            if sums.size > 0:
                local_max = sums.max()
                if local_max > max_sum:
                    max_sum = local_max
                    max_pos_in_sums = sums.argmax()
                    peak_index = shift + max_pos_in_sums
        if peak_index is not None:
            peak_indices = list(range(peak_index, peak_index + 4))
            peak_indices_list.append(peak_indices)
            trip_peak.append(dmodel['Time'].iloc[peak_index])
            trip_peak_vol.append(int(max_sum))
            max_single = results[peak_index:peak_index + 4].max()
            phf = max_sum / (max_single * 4) if max_single > 0 else 0
            trip_phf.append(round(phf, 3))
        else:
            trip_peak.append('N/A')
            trip_peak_vol.append(0)
            trip_phf.append(0)
            peak_indices_list.append([])
    return totals, day_totals, splits, trip_peak, trip_peak_vol, trip_phf, peak_indices_list


from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt

def generate_word_doc(trip_job, trip_client, trip_stationid, trip_location, trip_atr,
                      trip_site_des, trip_site_gps, dmodel, totals, day_totals, splits,
                      trip_peak, trip_peak_vol, trip_phf, peak_indices_list, first_date):
    # Generate a Word document with the data
    doc = Document()
    
    # Set document margins to narrow
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(36)  # Narrow margin (0.5 inches)
        section.bottom_margin = Pt(36)  # Narrow margin (0.5 inches)
        section.left_margin = Pt(36)  # Narrow margin (0.5 inches)
        section.right_margin = Pt(36)  # Narrow margin (0.5 inches)
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(8)
    
    # Add header paragraphs
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(f"Contract : {trip_job}\nClient: {trip_client}\nFacility: {trip_stationid}\nLocation: {trip_location}")
    p.paragraph_format.space_after = Pt(0)  # Set space after the paragraph in points
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f"ATR #: {trip_atr}\n{trip_site_des}\n{trip_site_gps}")
    p.paragraph_format.space_after = Pt(0)  # Set space after the paragraph in points
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("TRIP CONSULTANTS USA, INC.\n310 85th street, suite #A1\nBrooklyn, NY 11209\nTel (718) 833-6176 - Fax (718) 921-2844\n")
    p.paragraph_format.space_after = Pt(0)  # Set space after the paragraph in points

    # Increase the number of rows by 1 for the empty row
    num_rows = dmodel.shape[0] + 8  # Adding extra rows for totals, footers, and the empty row
    num_cols = dmodel.shape[1]
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    
    # First row adjustments
    first_row = table.rows[0]
    first_row.cells[0].text = "Start"  # Add "Start" in the very first cell
    texts = [first_date.strftime('%m/%d/%y')] + ['Tue', 'Wed', 'Thu', 'Fri', 'Sat', 'Sun', 'Average']
    text_index = 0
    for i in range(1, num_cols, 2):  # Start from the second cell (index 1)
        if text_index < len(texts): 
            cell_start = first_row.cells[i]
            cell_end = first_row.cells[i + 1] if i + 1 < num_cols else first_row.cells[i]
            merged_cell = cell_start.merge(cell_end)
            merged_cell.text = texts[text_index]
            merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            text_index += 1
    
    # Fill second row with text_1 to text_18
    second_row = table.rows[1]
    second_row_texts = ['Time'] + ['AM', 'PM'] * 8
    for idx, text in enumerate(second_row_texts):
        second_row.cells[idx].text = text
        second_row.cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add data rows with formatting
    for i in range(dmodel.shape[0]):
        row_cells = table.rows[i + 2].cells  # Adjust row index for the empty row
        for j, col_name in enumerate(dmodel.columns):
            cell = row_cells[j]
            cell.text = str(dmodel.iloc[i, j])
            # Apply formatting to peak hour cells
            if j > 0:  # Skip 'Time' column
                peak_indices = peak_indices_list[j - 1]
                if i in peak_indices:
                    # Apply background color
                    shading_elm = parse_xml(r'<w:shd {} w:fill="FCE6CE"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                    # Apply bold text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = False
                            run.font.size = Pt(8)
    
    # Helper function to center text and set font size
    def format_row_cells(row, font_size):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
    
    # Add totals row
    total_row_index = dmodel.shape[0] + 2
    total_row = table.rows[total_row_index]
    total_row.cells[0].text = 'Total'
    for idx, val in enumerate(totals):
        total_row.cells[idx + 1].text = str(int(val))
    format_row_cells(total_row, 6)
    
    # Add day totals row
    day_total_row_index = total_row_index + 1
    day_total_row = table.rows[day_total_row_index]
    day_total_row.cells[0].text = 'Day Total'
    for i in range(8):  # 7 days + average
        day_total = str(int(day_totals[i]))
        col_index = 1 + i * 2
        if col_index < num_cols:
            cell_start = day_total_row.cells[col_index]
            cell_end = day_total_row.cells[col_index + 1] if col_index + 1 < num_cols else cell_start
            merged_cell = cell_start.merge(cell_end)
            merged_cell.text = day_total
            merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_row_cells(day_total_row, 6)
    
    # Add splits row
    splits_row_index = day_total_row_index + 1
    splits_row = table.rows[splits_row_index]
    splits_row.cells[0].text = '% Splits'
    for idx, val in enumerate(splits):
        splits_row.cells[idx + 1].text = f"{val}%"
    format_row_cells(splits_row, 6)
    
    # Add peak info rows
    peak_row_index = splits_row_index + 1
    peak_row = table.rows[peak_row_index]
    peak_row.cells[0].text = 'Peak'
    for idx, val in enumerate(trip_peak):
        peak_row.cells[idx + 1].text = val if pd.notna(val) else 'N/A'
    format_row_cells(peak_row, 6)
    
    vol_row_index = peak_row_index + 1
    vol_row = table.rows[vol_row_index]
    vol_row.cells[0].text = 'Vol.'
    for idx, val in enumerate(trip_peak_vol):
        vol_row.cells[idx + 1].text = str(val)
    format_row_cells(vol_row, 6)
    
    phf_row_index = vol_row_index + 1
    phf_row = table.rows[phf_row_index]
    phf_row.cells[0].text = 'P.H.F.'
    for idx, val in enumerate(trip_phf):
        phf_row.cells[idx + 1].text = str(val)
    format_row_cells(phf_row, 6)
    
    # Add ADT and AADT
    p = doc.add_paragraph()
    ADT_val = int(sum(day_totals[:-1]) / 7)
    p.add_run(f"\nADT\tADT {ADT_val}\tAADT {ADT_val}")
    p.paragraph_format.space_after = Pt(0)  # Set space after the paragraph in points

    return doc


def generate_excel_file(trip_job, trip_client, trip_stationid, trip_location,
                        trip_atr, trip_site_des, trip_site_gps, dmodel, totals,
                        day_totals, splits, trip_peak, trip_peak_vol, trip_phf, first_date):
    # Generate an Excel file with the data
    wb = Workbook()
    ws = wb.active
    ws.title = 'Sheet1'
    
    # Write header information
    ws['A3'] = f"Contract : {trip_job}"
    ws['A4'] = f"Client: {trip_client}"
    ws['A5'] = f"Facility: {trip_stationid}"
    ws['A6'] = f"Location: {trip_location}"
    ws['N4'] = f"ATR #: {int(float(trip_atr)) if pd.notna(trip_atr) else 0}"  # Ensure integer formatting
    ws['N5'] = f"{trip_site_des}"
    ws['N6'] = f"{trip_site_gps}"
    
    # Merge cells D1 to M6 and fill with the required header text
    ws.merge_cells('D1:M6')
    merged_cell = ws['D1']
    merged_cell.value = (
        "TRIP CONSULTANTS USA, INC.\n"
        "310 85th Street, Suite #A1\n"
        "Brooklyn, NY 11209\n"
        "Tel (718) 833-6176 - Fax (718) 921-2844"
    )
    merged_cell.font = Font(bold=False, name='MS Sans Serif', size=8)
    merged_cell.alignment = Alignment(horizontal='center', vertical='top', wrap_text=True)
    
    # Add the day-of-the-week row above "A.M." and "P.M."
    days = [first_date.strftime('%m/%d/%y')] + ['Tue', 'Wed', 'Thu', 'Fri', 'Sat', 'Sun', 'Average Day']
    for idx, day in enumerate(days):
        col_idx = 2 + idx * 2
        ws.merge_cells(start_row=10, start_column=col_idx, end_row=10, end_column=col_idx + 1)
        cell = ws.cell(row=10, column=col_idx)
        cell.value = day
        cell.font = Font(bold=False, name='MS Sans Serif', size=8)
        cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # Write the data table headers (A.M. and P.M.)
    for idx in range(len(dmodel.columns)-1):
        cell = ws.cell(row=11, column=idx + 2)
        cell.value = "A.M." if idx % 2 == 0 else "P.M."
        cell.font = Font(bold=False, name='MS Sans Serif', size=8)
        cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # Add data rows with formatting
    for i in range(dmodel.shape[0]):
        for j in range(dmodel.shape[1]):
            cell = ws.cell(row=12 + i, column=1 + j)
            if j == 0:  # Time column
                time_value = dmodel.iloc[i, j].strip()
                cell.value = time_value.replace('0:00', '12:00').replace('0:15', '12:15').replace('0:30', '12:30').replace('0:45', '12:45')
            else:
                cell.value = dmodel.iloc[i, j]
            cell.font = Font(name='MS Sans Serif', size=8)
            cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # Alternate row coloring from row 12 (after header row)
    for i in range(dmodel.shape[0]):
        row_number = 12 + i
        fill_color = 'B3B3B3' if i % 2 == 0 else 'FFFFFF'  # Grey and white
        for col in range(1, dmodel.shape[1] + 1):
            ws.cell(row=row_number, column=col).fill = PatternFill(start_color=fill_color, end_color=fill_color, fill_type='solid')
    
    # Calculate the indices for the additional rows
    total_row = 12 + dmodel.shape[0]
    day_total_row = total_row + 1
    splits_row = day_total_row + 1
    empty_row_1 = splits_row + 1
    empty_row_2 = empty_row_1 + 1
    peak_row = empty_row_2 + 1
    vol_row = peak_row + 1
    phf_row = vol_row + 1
    
    # Write totals row
    ws.cell(row=total_row, column=1).value = 'Total'
    ws.cell(row=total_row, column=1).font = Font(name='MS Sans Serif', size=8)
    for idx, val in enumerate(totals):
        cell = ws.cell(row=total_row, column=2+idx)
        cell.value = int(val)
        cell.font = Font(name='MS Sans Serif', size=8)
        cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # Write day totals row with "Day Total" label
    ws.cell(row=day_total_row, column=1).value = 'Day Total'  # Add "Day Total" in the first column
    ws.cell(row=day_total_row, column=1).font = Font(name='MS Sans Serif', size=8)
    for i, day_total in enumerate(day_totals):
        col_index = 2 + i * 2
        ws.merge_cells(start_row=day_total_row, start_column=col_index, end_row=day_total_row, end_column=col_index + 1)
        cell = ws.cell(row=day_total_row, column=col_index)
        cell.value = int(day_total)
        cell.font = Font(name='MS Sans Serif', size=8)
        cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # Write splits row
    ws.cell(row=splits_row, column=1).value = '% Splits'
    ws.cell(row=splits_row, column=1).font = Font(name='MS Sans Serif', size=8)
    for idx, val in enumerate(splits):
        cell = ws.cell(row=splits_row, column=2+idx)
        cell.value = f"{val}%"
        cell.font = Font(name='MS Sans Serif', size=8)
        cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # Write peak row
    ws.cell(row=peak_row, column=1).value = 'Peak'
    ws.cell(row=peak_row, column=1).font = Font(name='MS Sans Serif', size=8)
    for idx, val in enumerate(trip_peak):
        cell = ws.cell(row=peak_row, column=2+idx)
        cell.value = val if pd.notna(val) else 'N/A'
        cell.font = Font(name='MS Sans Serif', size=8)
        cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # Write vol row
    ws.cell(row=vol_row, column=1).value = 'Vol.'
    ws.cell(row=vol_row, column=1).font = Font(name='MS Sans Serif', size=8)
    for idx, val in enumerate(trip_peak_vol):
        cell = ws.cell(row=vol_row, column=2+idx)
        cell.value = val
        cell.font = Font(name='MS Sans Serif', size=8)
        cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # Write PHF row
    ws.cell(row=phf_row, column=1).value = 'P.H.F.'
    ws.cell(row=phf_row, column=1).font = Font(name='MS Sans Serif', size=8)
    for idx, val in enumerate(trip_phf):
        cell = ws.cell(row=phf_row, column=2+idx)
        cell.value = val
        cell.font = Font(name='MS Sans Serif', size=8)
        cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # Set column widths
    for col in range(1, 18):  # Adjust the number of columns as needed
        ws.column_dimensions[chr(64 + col)].width = 6  # A=1, B=2...
    
    # Set row heights
    for row in range(1, ws.max_row + 1):
        ws.row_dimensions[row].height = 12
    
    return wb


def generate_zip(result_list):
    # Generate a zip file containing the Word documents and Excel files
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
        for loc_name, data in result_list.items():
            # Save Word document to bytes
            doc_buffer = io.BytesIO()
            data['doc'].save(doc_buffer)
            zip_file.writestr(f"{loc_name}.docx", doc_buffer.getvalue())
            # Save Excel workbook to bytes
            excel_buffer = io.BytesIO()
            data['excel'].save(excel_buffer)
            zip_file.writestr(f"{loc_name}.xlsx", excel_buffer.getvalue())
    zip_buffer.seek(0)
    return zip_buffer

if __name__ == '__main__':
    main()
